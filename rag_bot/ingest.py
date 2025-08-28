import os, re, pickle, json, requests
from typing import List, Dict
import streamlit as st
import numpy as np
from PyPDF2 import PdfReader
import pdfplumber
from docx import Document
from tqdm import tqdm
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
import faiss
from dotenv import load_dotenv

load_dotenv()

DOCS_DIR = "docs"
INDEX_DIR = "index"
EMBED_MODEL_NAME = os.getenv('EMBED_MODEL_1')

os.makedirs(INDEX_DIR, exist_ok=True)

def read_pdf(path: str) -> str:
    texts = []
    with pdfplumber.open(path) as pdf:
        for p in pdf.pages:
            t = p.extract_text() or ""
            texts.append(t)
    return "\n".join(texts)

def read_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

def load_documents(folder: str) -> List[Dict]:
    items = []
    for fn in os.listdir(folder):
        fp = os.path.join(folder, fn)
        if fn.lower().endswith(".pdf"):
            text = read_pdf(fp)
        elif fn.lower().endswith(".docx"):
            text = read_docx(fp)
        else:
            continue
        text = re.sub(r"\s+\n", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        items.append({"source": fn, "text": text})
    return items

def split_sections(text: str) -> Dict[str, str]:
    # repère des sections financières FR/EN courantes
    def extract_text_docx(file):
        doc = Document(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    pattern = r"(Bilan|Balance Sheet|Compte de résultat|Income Statement|Cash[- ]?flow|Flux de trésorerie|Annexes|Notes)"
    matches = list(re.finditer(pattern, text, re.IGNORECASE))
    if not matches:
        return {"Document": text}
    sections = {}
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(text)
        name = m.group().strip().title()
        sections[name] = text[start:end].strip()
    return sections

def build_index():
    docs = load_documents(DOCS_DIR)
    if not docs:
        raise SystemExit("Aucun document trouvé dans ./docs")

    # Interface Streamlit pour chat RAG sur document uploadé
    def main_streamlit():


        API_KEY = os.getenv("API_GROQ")
        GROQ_API_URL = "https://api.groqcloud.com/v1/chat/completions"

        def extract_text_pdf(file):
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

        def extract_text_docx(file):
            doc = Document(file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text

        def ask_groqcloud(question, context):
            headers = {
                "Authorization": f"Bearer {API_KEY}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "mixtral-8x7b-32768",  # ou le modèle Groq que tu utilises
                "messages": [
                    {"role": "system", "content": "Tu es un assistant qui répond uniquement en te basant sur le contexte fourni."},
                    {"role": "user", "content": f"Contexte:\n{context}\n\nQuestion: {question}"}
                ]
            }
            response = requests.post(GROQ_API_URL, headers=headers, json=payload)
            if response.status_code == 200:
                return response.json()["choices"][0]["message"]["content"]
            else:
                return f"Erreur API: {response.status_code} - {response.text}"

        st.title("Chatbot - RAG")
        uploaded_file = st.file_uploader("Uploader un document (PDF ou DOCX)", type=["pdf", "docx"])

        document_text = ""
        if uploaded_file:
            if uploaded_file.type == "application/pdf":
                document_text = extract_text_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                document_text = extract_text_docx(uploaded_file)
            st.success("Document chargé et texte extrait !")
            st.text_area("Contenu extrait", document_text, height=200)

        if document_text:
            question = st.text_input("Posez votre question sur le document :")
            if question:
                with st.spinner("Recherche de la réponse..."):
                    answer = ask_groqcloud(question, document_text)
                st.markdown(f"**Réponse :** {answer}")

    if __name__ == "__main__":
        main_streamlit()

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)

    all_chunks: List[str] = []
    metas: List[Dict] = []

    for d in docs:
        sections = split_sections(d["text"])
        for sec_name, sec_text in sections.items():
            if not sec_text.strip():
                continue
            parts = splitter.split_text(sec_text)
            for j, ch in enumerate(parts):
                all_chunks.append(ch)
                metas.append({"source": d["source"], "section": sec_name, "chunk_id": j})

    print(f"Total chunks: {len(all_chunks)}")

    model = SentenceTransformer(EMBED_MODEL_NAME)
    emb = model.encode(all_chunks, batch_size=64, convert_to_numpy=True, show_progress_bar=True)
    # normalisation + similarité cosinus via produit scalaire
    norms = np.linalg.norm(emb, axis=1, keepdims=True) + 1e-12
    emb = emb / norms

    dim = emb.shape[1]
    index = faiss.IndexFlatIP(dim)  # produit scalaire
    index.add(emb.astype(np.float32))

    faiss.write_index(index, os.path.join(INDEX_DIR, "faiss.index"))
    with open(os.path.join(INDEX_DIR, "chunks.pkl"), "wb") as f:
        pickle.dump(all_chunks, f)
    with open(os.path.join(INDEX_DIR, "metas.pkl"), "wb") as f:
        pickle.dump(metas, f)

    cfg = {"embed_model": EMBED_MODEL_NAME}
    with open(os.path.join(INDEX_DIR, "config.json"), "w") as f:
        json.dump(cfg, f, indent=2)

    print("Index construit et sauvegardé dans ./index")

if __name__ == "__main__":
    build_index()
